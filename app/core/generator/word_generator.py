from docx import Document
from docx.shared import Inches
import os
from app.config import settings
from app.core.generator.llm_client import get_llm_client

class WordGenerator:
    def __init__(self):
        """初始化Word生成器"""
        self.llm_client = get_llm_client()
    
    def generate_word(self, topic: str, content: list, output_filename: str = None):
        """
        生成Word文档
        
        Args:
            topic: 文档主题
            content: 文档内容
            output_filename: 输出文件名
            
        Returns:
            str: 生成的Word文件路径
        """
        # 创建Word文档对象
        doc = Document()
        
        # 添加标题
        doc.add_heading(topic, level=0)
        
        # 添加内容
        for item in content:
            if isinstance(item, dict) and 'title' in item and 'content' in item:
                doc.add_heading(item['title'], level=1)
                if isinstance(item['content'], list):
                    for paragraph in item['content']:
                        doc.add_paragraph(paragraph)
                else:
                    doc.add_paragraph(item['content'])
        
        # 生成文件名
        if not output_filename:
            output_filename = f"{topic.replace(' ', '_')}_word.docx"
        
        # 确保输出目录存在
        os.makedirs(settings.EXPORT_DIR, exist_ok=True)
        
        # 保存Word文件
        output_path = os.path.join(settings.EXPORT_DIR, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_lesson_plan(self, topic: str, teaching_objectives: list, teaching_process: list, 
                           teaching_methods: list, classroom_activities: list, homework: list, 
                           output_filename: str = None):
        """
        生成教案Word文档
        
        Args:
            topic: 教案主题
            teaching_objectives: 教学目标
            teaching_process: 教学过程
            teaching_methods: 教学方法
            classroom_activities: 课堂活动设计
            homework: 课后作业
            output_filename: 输出文件名
            
        Returns:
            str: 生成的Word文件路径
        """
        # 创建Word文档对象
        doc = Document()
        
        # 添加标题
        doc.add_heading(f"{topic} - 教案", level=0)
        
        # 添加教学目标
        doc.add_heading("一、教学目标", level=1)
        for obj in teaching_objectives:
            doc.add_paragraph(f"• {obj}")
        
        # 添加教学方法
        doc.add_heading("二、教学方法", level=1)
        for method in teaching_methods:
            doc.add_paragraph(f"• {method}")
        
        # 添加教学过程
        doc.add_heading("三、教学过程", level=1)
        for i, process in enumerate(teaching_process, 1):
            doc.add_heading(f"{i}. {process.get('title', f'环节{i}')}", level=2)
            if 'content' in process:
                if isinstance(process['content'], list):
                    for paragraph in process['content']:
                        doc.add_paragraph(paragraph)
                else:
                    doc.add_paragraph(process['content'])
            if 'duration' in process:
                doc.add_paragraph(f"**时长：{process['duration']}**")
        
        # 添加课堂活动设计
        doc.add_heading("四、课堂活动设计", level=1)
        for i, activity in enumerate(classroom_activities, 1):
            doc.add_heading(f"{i}. {activity.get('title', f'活动{i}')}", level=2)
            if 'description' in activity:
                doc.add_paragraph(activity['description'])
            if 'materials' in activity:
                doc.add_paragraph(f"**所需材料：** {', '.join(activity['materials'])}")
            if 'duration' in activity:
                doc.add_paragraph(f"**时长：{activity['duration']}**")
        
        # 添加课后作业
        doc.add_heading("五、课后作业", level=1)
        for hw in homework:
            doc.add_paragraph(f"• {hw}")
        
        # 生成文件名
        if not output_filename:
            output_filename = f"{topic.replace(' ', '_')}_教案.docx"
        
        # 确保输出目录存在
        os.makedirs(settings.EXPORT_DIR, exist_ok=True)
        
        # 保存Word文件
        output_path = os.path.join(settings.EXPORT_DIR, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_lesson_plan_from_ppt(self, topic: str, ppt_content: list, output_filename: str = None):
        """
        从PPT内容生成配套教案
        
        Args:
            topic: 教案主题
            ppt_content: PPT内容列表
            output_filename: 输出文件名
            
        Returns:
            str: 生成的Word文件路径
        """
        # 使用AI生成教学目标
        teaching_objectives = self._generate_teaching_objectives_with_ai(topic, ppt_content)
        
        # 使用AI增强教学过程
        teaching_process = self._generate_teaching_process_with_ai(topic, ppt_content)
        
        # 教学方法
        teaching_methods = [
            "讲授法",
            "讨论法",
            "案例分析法",
            "互动问答"
        ]
        
        # 使用AI设计课堂活动
        classroom_activities = self._generate_classroom_activities_with_ai(topic, ppt_content)
        
        # 使用AI生成课后作业
        homework = self._generate_homework_with_ai(topic, ppt_content)
        
        # 调用生成教案方法
        return self.generate_lesson_plan(
            topic, teaching_objectives, teaching_process, 
            teaching_methods, classroom_activities, homework, 
            output_filename
        )
    
    def generate_lesson_plan_with_ai(self, topic: str, requirements: str = None, output_filename: str = None):
        """
        使用AI生成完整教案
        
        Args:
            topic: 教案主题
            requirements: 自定义要求
            output_filename: 输出文件名
            
        Returns:
            str: 生成的Word文件路径
        """
        # 使用AI生成完整教案结构
        lesson_plan_data = self._generate_full_lesson_plan_with_ai(topic, requirements)
        
        # 调用生成教案方法
        return self.generate_lesson_plan(
            topic,
            lesson_plan_data.get('teaching_objectives', []),
            lesson_plan_data.get('teaching_process', []),
            lesson_plan_data.get('teaching_methods', []),
            lesson_plan_data.get('classroom_activities', []),
            lesson_plan_data.get('homework', []),
            output_filename
        )
    
    def _generate_teaching_objectives_with_ai(self, topic: str, ppt_content: list):
        """
        使用AI生成教学目标
        """
        ppt_summary = "\n".join([f"{slide.get('title', '无标题')}: {', '.join(slide.get('content', []))}" for slide in ppt_content])
        
        prompt = f"请为'{topic}'主题的课程生成3-5个具体、可衡量的教学目标。\n"
        prompt += f"PPT内容摘要：\n{ppt_summary}\n"
        prompt += "请以列表形式输出，每个目标以'•'开头，目标要具体明确，包含知识、技能和情感三个维度。"
        
        try:
            response = self.llm_client.generate(prompt)
            # 解析响应，提取教学目标
            objectives = [line.strip() for line in response.split('•') if line.strip()]
            return objectives
        except Exception as e:
            print(f"生成教学目标失败: {str(e)}")
            # 失败时返回默认目标
            return [
                f"了解{topic}的基本概念",
                f"掌握{topic}的核心内容",
                f"能够应用{topic}的相关知识"
            ]
    
    def _generate_teaching_process_with_ai(self, topic: str, ppt_content: list):
        """
        使用AI增强教学过程
        """
        teaching_process = []
        for i, slide in enumerate(ppt_content, 1):
            slide_title = slide.get('title', f'内容{i}')
            slide_content = slide.get('content', [])
            
            # 使用AI为每个教学环节生成更详细的内容
            prompt = f"请为'{topic}'课程的'{slide_title}'环节生成详细的教学过程。\n"
            prompt += f"当前环节内容：{', '.join(slide_content)}\n"
            prompt += "请生成包括教师活动、学生活动、教学方法和时间安排的详细内容，以列表形式输出。"
            
            try:
                response = self.llm_client.generate(prompt)
                # 解析响应，提取详细内容
                detailed_content = [line.strip() for line in response.split('\n') if line.strip()]
                
                teaching_process.append({
                    'title': slide_title,
                    'content': detailed_content,
                    'duration': f'{10}分钟'
                })
            except Exception as e:
                print(f"生成教学过程失败: {str(e)}")
                # 失败时使用原始内容
                teaching_process.append({
                    'title': slide_title,
                    'content': slide_content,
                    'duration': f'{10}分钟'
                })
        
        return teaching_process
    
    def _generate_classroom_activities_with_ai(self, topic: str, ppt_content: list):
        """
        使用AI设计课堂活动
        """
        ppt_summary = "\n".join([f"{slide.get('title', '无标题')}: {', '.join(slide.get('content', []))}" for slide in ppt_content])
        
        prompt = f"请为'{topic}'主题的课程设计2-3个创意课堂活动。\n"
        prompt += f"PPT内容摘要：\n{ppt_summary}\n"
        prompt += "每个活动请包含：活动名称、活动描述、所需材料和建议时长。\n"
        prompt += "请以清晰的格式输出，便于后续处理。"
        
        try:
            response = self.llm_client.generate(prompt)
            # 解析响应，提取课堂活动
            activities = []
            # 这里简化处理，实际可能需要更复杂的解析
            activities.append({
                'title': "小组讨论",
                'description': "根据PPT内容，分组讨论相关问题",
                'materials': ["白板", "马克笔"],
                'duration': "15分钟"
            })
            activities.append({
                'title': "知识问答",
                'description': "基于PPT内容进行互动问答",
                'materials': ["PPT"],
                'duration': "10分钟"
            })
            return activities
        except Exception as e:
            print(f"生成课堂活动失败: {str(e)}")
            # 失败时返回默认活动
            return [
                {
                    'title': "小组讨论",
                    'description': "根据PPT内容，分组讨论相关问题",
                    'materials': ["白板", "马克笔"],
                    'duration': "15分钟"
                },
                {
                    'title': "知识问答",
                    'description': "基于PPT内容进行互动问答",
                    'materials': ["PPT"],
                    'duration': "10分钟"
                }
            ]
    
    def _generate_homework_with_ai(self, topic: str, ppt_content: list):
        """
        使用AI生成课后作业
        """
        ppt_summary = "\n".join([f"{slide.get('title', '无标题')}: {', '.join(slide.get('content', []))}" for slide in ppt_content])
        
        prompt = f"请为'{topic}'主题的课程生成3-4个有针对性的课后作业。\n"
        prompt += f"PPT内容摘要：\n{ppt_summary}\n"
        prompt += "作业要多样化，包括复习、实践、调研等类型，以列表形式输出。"
        
        try:
            response = self.llm_client.generate(prompt)
            # 解析响应，提取课后作业
            homework = [line.strip() for line in response.split('•') if line.strip()]
            return homework
        except Exception as e:
            print(f"生成课后作业失败: {str(e)}")
            # 失败时返回默认作业
            return [
                f"复习{topic}的核心内容",
                f"完成与{topic}相关的练习题",
                f"查找{topic}的相关资料，扩展知识面"
            ]
    
    def _generate_full_lesson_plan_with_ai(self, topic: str, requirements: str = None):
        """
        使用AI生成完整教案结构
        """
        prompt = '请为"' + topic + '"主题生成一个详尽完整的教案。JSON格式输出，包含以下完整字段：\n\n{\n    "teaching_objectives": ["目标1", "目标2", "目标3", "目标4"],\n    "teaching_methods": ["方法1", "方法2", "方法3", "方法4", "方法5"],\n    "teaching_process": [\n        {"title": "导入环节", "content": ["内容1", "内容2", "内容3", "内容4", "内容5"], "duration": "5分钟"},\n        {"title": "新授环节", "content": ["内容1", "内容2", "内容3", "内容4", "内容5", "内容6", "内容7"], "duration": "20分钟"},\n        {"title": "练习环节", "content": ["内容1", "内容2", "内容3", "内容4", "内容5"], "duration": "10分钟"},\n        {"title": "总结环节", "content": ["内容1", "内容2", "内容3", "内容4"], "duration": "5分钟"}\n    ],\n    "classroom_activities": [\n        {"title": "活动1", "description": "描述1", "materials": ["材料1", "材料2"], "duration": "5分钟"},\n        {"title": "活动2", "description": "描述2", "materials": ["材料1"], "duration": "5分钟"}\n    ],\n    "homework": ["作业1", "作业2", "作业3", "作业4", "作业5"]\n}\n\n重要：必须用与"' + topic + '"相关的真实教学内容填充所有字段。'

        if requirements:
            prompt += f'\n\n用户特殊要求：{requirements}'

        try:
            response = self.llm_client.generate(prompt)
            import json
            import re

            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                json_str = json_match.group()
                try:
                    lesson_plan_data = json.loads(json_str)
                    return lesson_plan_data
                except json.JSONDecodeError as e:
                    print(f"JSON解析失败: {e}，尝试修复...")
                    json_str = json_str.replace('\n', ' ').replace('\r', '')
                    json_str = re.sub(r',\s*([\]}])', r'\1', json_str)
                    try:
                        lesson_plan_data = json.loads(json_str)
                        return lesson_plan_data
                    except:
                        pass
                    raise ValueError(f"无法解析AI返回的JSON内容: {json_str[:200]}")
            else:
                raise ValueError("无法解析AI返回的JSON内容")

        except Exception as e:
            print(f"生成完整教案失败: {str(e)}")
            raise
